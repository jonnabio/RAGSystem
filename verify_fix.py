import os
import httpx
from docx import Document
import asyncio

async def verify_fix():
    # 1. Create a large docx file to ensure > 100 chunks
    # Chunk size is 512, overlap 50.
    # To get > 100 chunks, we need ~ 50,000 characters.
    doc = Document()
    doc.add_heading('Stress Test Document', 0)

    paragraph_text = "This is a stress test paragraph for the RAG system. " * 20
    for i in range(200): # 200 paragraphs
        doc.add_paragraph(f"Paragraph {i}: {paragraph_text}")

    test_file = "stress_test.docx"
    doc.save(test_file)
    print(f"Created {test_file}")

    # 2. Upload the document
    url = "http://localhost:8001/api/documents/upload"
    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            with open(test_file, "rb") as f:
                files = {"file": (test_file, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
                print(f"Uploading {test_file} to {url}...")
                response = await client.post(url, files=files)

            if response.status_code == 200:
                print("✅ Upload successful!")
                print(response.json())
            else:
                print(f"❌ Upload failed with status {response.status_code}")
                print(response.text)
    except Exception as e:
        print(f"❌ Error during upload: {e}")
    finally:
        if os.path.exists(test_file):
            os.remove(test_file)

if __name__ == "__main__":
    asyncio.run(verify_fix())
