"""Word document processor using python-docx."""

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from typing import Dict, Any
from pathlib import Path
import logging

logger = logging.getLogger(__name__)


class WordProcessor:
    """
    Word document processor for DOCX files.

    Uses python-docx library for document parsing.
    """

    def extract_text(self, file_path: str) -> str:
        """
        Extract all text from Word document.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content

        Raises:
            FileNotFoundError: If file doesn't exist
            RuntimeError: If document is corrupted or not a valid DOCX
        """
        if not Path(file_path).exists():
            raise FileNotFoundError(f"Word file not found: {file_path}")

        try:
            doc = DocxDocument(file_path)

            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)

            if not paragraphs:
                logger.warning(f"No text extracted from {file_path}")
                return ""

            return "\n\n".join(paragraphs)

        except PackageNotFoundError:
            raise RuntimeError(f"Invalid or corrupted DOCX file: {file_path}")
        except Exception as e:
            raise RuntimeError(f"Error processing Word document: {e}")

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract Word document metadata.

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary containing document metadata
        """
        if not Path(file_path).exists():
            raise FileNotFoundError(f"Word file not found: {file_path}")

        try:
            doc = DocxDocument(file_path)
            core_props = doc.core_properties

            metadata = {
                "format": "DOCX",
                "author": core_props.author or "",
                "title": core_props.title or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "creationDate": core_props.created.isoformat() if core_props.created else "",
                "modDate": core_props.modified.isoformat() if core_props.modified else "",
                "creator": "python-docx",
                "producer": "python-docx",
                "page_count": 0,
                "word_count": sum(len(p.text.split()) for p in doc.paragraphs)
            }

            return metadata

        except Exception as e:
            logger.error(f"Error extracting Word metadata: {e}")
            return {"format": "DOCX", "error": str(e)}

    def get_paragraph_count(self, file_path: str) -> int:
        """
        Get number of paragraphs in document.

        Args:
            file_path: Path to DOCX file

        Returns:
            Number of paragraphs
        """
        try:
            doc = DocxDocument(file_path)
            return len(doc.paragraphs)
        except Exception as e:
            logger.error(f"Error getting paragraph count: {e}")
            return 0
